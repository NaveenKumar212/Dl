import streamlit as st
import os
import re
from io import BytesIO
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.question_answering import load_qa_chain
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.llms import OpenAI
from langchain_community.callbacks.manager import get_openai_callback
from langchain.schema import Document
from docx import Document as DocxDocument
import pandas as pd
import pytesseract
from PIL import Image
import logging
import hashlib
import json
import base64
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import pickle

# Place st.set_page_config() at the very beginning
st.set_page_config(page_title="User Library", layout="wide")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Streamlit Cloud Storage Configuration
USER_DATA_FILE = "user_data.json"
UPLOADED_FILES_DIR = "uploaded_files"
VECTOR_INDEX_DIR = "vector_indices"

# Create necessary directories
os.makedirs(UPLOADED_FILES_DIR, exist_ok=True)
os.makedirs(VECTOR_INDEX_DIR, exist_ok=True)

# Get OpenAI API key from Streamlit secrets
try:
    OPENAI_API_KEY = st.secrets["openai"]["api_key"]
    if not OPENAI_API_KEY or OPENAI_API_KEY.startswith("sk-..."):
        st.error("Please set your OpenAI API key in the .streamlit/secrets.toml file")
        st.stop()
except Exception as e:
    st.error("Error loading OpenAI API key. Please check your secrets.toml file")
    st.stop()

# Initialize the sentence transformer model
@st.cache_resource(show_spinner=False)
def get_embedding_model():
    try:
        model = SentenceTransformer('all-MiniLM-L6-v2')
        return model
    except Exception as e:
        st.error(f"Error loading embedding model: {str(e)}")
        return None

def load_user_data():
    """Load user data from Streamlit's session state."""
    if 'user_data' not in st.session_state:
        try:
            with open(USER_DATA_FILE, 'r') as f:
                data = json.load(f)
                st.session_state.user_data = data
                logging.info("User data loaded successfully.")
                return data
        except FileNotFoundError:
            logging.info("User data file not found. Returning empty dictionary.")
            st.session_state.user_data = {}
            return {}
    return st.session_state.user_data

def save_user_data(user_data):
    """Save user data to Streamlit's session state and local file."""
    st.session_state.user_data = user_data
    with open(USER_DATA_FILE, 'w') as f:
        json.dump(user_data, f)
    logging.info("User data saved successfully.")

def save_uploaded_file(file, username):
    """Save an uploaded file permanently to the filesystem."""
    file_name = f"{username}_{file.name}"
    file_path = os.path.join(UPLOADED_FILES_DIR, file_name)
    
    # Read file contents
    file_contents = file.read()
    
    # Save file to disk
    with open(file_path, 'wb') as f:
        f.write(file_contents)
    
    return file_name

def get_uploaded_file(file_name):
    """Get an uploaded file from the filesystem."""
    file_path = os.path.join(UPLOADED_FILES_DIR, file_name)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as f:
            return BytesIO(f.read())
    return None

def delete_uploaded_file(file_name):
    """Delete an uploaded file from the filesystem."""
    file_path = os.path.join(UPLOADED_FILES_DIR, file_name)
    if os.path.exists(file_path):
        os.remove(file_path)

def delete_all_uploaded_files(username):
    """Delete all uploaded files for a user from the filesystem."""
    for file_name in os.listdir(UPLOADED_FILES_DIR):
        if file_name.startswith(f"{username}_"):
            file_path = os.path.join(UPLOADED_FILES_DIR, file_name)
            if os.path.exists(file_path):
                os.remove(file_path)

def hash_password(password):
    """Hash the provided password."""
    return hashlib.sha256(password.encode()).hexdigest()

hide_st_style = """
    <style>
    header {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
    """
st.markdown(hide_st_style, unsafe_allow_html=True)

def check_password(username, password, user_data):
    """Check if the provided password matches the stored hash."""
    hashed_password = hash_password(password)
    return user_data.get(username, {}).get('password') == hashed_password

def signup():
    """Display signup form and handle user registration."""
    st.title("Signup")
    new_username = st.text_input("New Username")
    new_email = st.text_input("Email")
    new_password = st.text_input("New Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")

    if st.button("Signup"):
        user_data = load_user_data()
        if new_username in user_data:
            st.error("Username already exists.")
        elif new_password != confirm_password:
            st.error("Passwords do not match.")
        else:
            user_data[new_username] = {
                'password': hash_password(new_password),
                'email': new_email,
                'subscription': 'free',  # Default to free subscription
                'uploaded_files': [],  # Add user-specific file tracking
                'faiss_index_path': f"faiss_index_{new_username}"  # User Specific FAISS index path.
            }
            save_user_data(user_data)
            st.success("Signup successful! Please log in.")
            st.query_params["page"] = "login"
            st.rerun()
    st.markdown("If You have account [Login](?page=login)")

def login():
    """Display login form and handle authentication."""
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        user_data = load_user_data()
        if check_password(username, password, user_data):
            st.query_params["page"] = "main"
            st.query_params["username"] = username
            st.rerun()
        else:
            st.error("Incorrect username or password.")
    st.markdown("If You Don't have account [Signup](?page=signup)")

def upgrade_subscription(selected_subscription, username):
    """Upgrade subscription logic."""
    if not username:
        st.error("You need to be logged in to upgrade your subscription.")
        return

    user_data = load_user_data()

    if username not in user_data:
        st.error("User not found in database!")
        return

    try:
        user_data = load_user_data()  # Reload user data
        logging.info(f"Before update: {user_data.get(username)}")

        user_data[username]['subscription'] = selected_subscription.lower()
        save_user_data(user_data)

        logging.info(f"After update: {user_data.get(username)}")

        st.success(f"Subscription upgraded to {selected_subscription}!")
    except Exception as e:
        st.error(f"An error occurred during upgrade: {e}")
        logging.error(f"Subscription upgrade error: {e}")

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file."""
    doc = DocxDocument(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_contacts(text):
    phone_numbers = re.findall(r'\+?\d{1,3}[\s.-]?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}', text)
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)
    return phone_numbers, emails

def store_embeddings_permanently(all_chunks, username):
    """Store document embeddings using sentence-transformers and FAISS."""
    try:
        model = get_embedding_model()
        if model is None:
            st.error("Failed to load embedding model")
            return
            
        documents = [Document(page_content=chunk) for chunk in all_chunks]
        
        # Generate embeddings for all chunks
        texts = [doc.page_content for doc in documents]
        embeddings = model.encode(texts, show_progress_bar=True)
        
        # Create FAISS index
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings.astype('float32'))
        
        # Save the index and documents
        index_path = os.path.join(VECTOR_INDEX_DIR, f"{username}_index")
        faiss.write_index(index, f"{index_path}.faiss")
        with open(f"{index_path}_documents.pkl", 'wb') as f:
            pickle.dump(documents, f)
        
        logging.info(f"Vector index stored successfully for user: {username}")
    except Exception as e:
        logging.error(f"Error storing vector index: {e}")
        st.error(f"Error storing embeddings: {str(e)}")

def load_stored_embeddings(username):
    """Load stored vector embeddings if available."""
    try:
        index_path = os.path.join(VECTOR_INDEX_DIR, f"{username}_index")
        
        if not os.path.exists(f"{index_path}.faiss") or not os.path.exists(f"{index_path}_documents.pkl"):
            logging.info("No existing vector index found.")
            return None, None
        
        # Load the index and documents
        index = faiss.read_index(f"{index_path}.faiss")
        with open(f"{index_path}_documents.pkl", 'rb') as f:
            documents = pickle.load(f)
        
        logging.info("Vector index loaded successfully.")
        return index, documents
    except Exception as e:
        logging.error(f"Error loading vector index: {e}")
        return None, None

def extract_text_from_image(image_file):
    """Extract text from an image file using Tesseract OCR."""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting text from image: {e}")
        logging.error(f"Image Extraction error: {e}")
        return ""

def change_user_subscription(username, new_subscription):
    """Changes the subscription status of a user."""
    user_data = load_user_data()  # Load the current user data

    if username in user_data:
        user_data[username]['subscription'] = new_subscription.lower()  # Update the subscription
        save_user_data(user_data)  # Save the updated data
        logging.info(f"Subscription for {username} changed to {new_subscription}")
        return True
    else:
        logging.error(f"User {username} not found.")
        return False

def admin_page(username):
    """Admin interface to manage user subscriptions."""
    st.title("Admin - Manage Subscriptions")

    user_data = load_user_data()
    usernames = list(user_data.keys())

    if not usernames:
        st.info("No users found.")
        return

    selected_user = st.selectbox("Select User", usernames, key="admin_user_select")
    current_subscription = user_data[selected_user]['subscription']
    st.write(f"Current Subscription: {current_subscription.capitalize()}")

    subscription_options = ["free", "silver", "gold", "platinum"] # All lower case
    selected_subscription = st.selectbox("New Subscription", subscription_options)

    if st.button("Change Subscription"):
        if change_user_subscription(selected_user, selected_subscription):
            st.success(f"Subscription for {selected_user} changed to {selected_subscription}!")
            st.rerun()  # refresh the page
        else:
            st.error(f"Failed to change subscription for {selected_user}.")

    # Convert subscription to link
    st.markdown("---")
    st.subheader("User Subscriptions")
    for user, data in user_data.items():
        subscription = data['subscription'].capitalize()
        st.markdown(f"**{user}**: [{subscription}](?page=admin&user={user}&username={username})")

    # Handle query parameters for direct subscription change
    if "user" in st.query_params:
        user_to_change = st.query_params["user"]
        if user_to_change in user_data:
            st.markdown("---")
            st.subheader(f"Change Subscription for {user_to_change}")
            current_subscription_query = user_data[user_to_change]['subscription'].capitalize()
            st.write(f"Current Subscription: {current_subscription_query}")

            subscription_options_query = ["free", "silver", "gold", "platinum"] # All lower case
            selected_subscription_query = st.selectbox("New Subscription", subscription_options_query, key=f"query_{user_to_change}")

            if st.button("Change Subscription", key=f"change_query_{user_to_change}"):
                if change_user_subscription(user_to_change, selected_subscription_query):
                    st.success(f"Subscription for {user_to_change} changed to {selected_subscription_query}!")
                    st.query_params["page"] = "admin"
                    st.query_params["username"] = username
                    st.rerun()
                else:
                    st.error(f"Failed to change subscription for {user_to_change}.")
        else:
            st.error(f"User {user_to_change} not found.")

def subscription_model_page():
    """Display the subscription model information."""
    st.title("Subscription Models")
    st.write("""
    Here's an overview of our subscription models:
    """)

    st.subheader("Free")
    st.write("- Limited to 15 uploaded files.")
    st.write("- Basic features.")

    st.subheader("Silver")
    st.write("- Up to 50 uploaded files.")
    st.write("- Enhanced features.")

    st.subheader("Gold")
    st.write("- Up to 150 uploaded files.")
    st.write("- Advanced features.")

    st.subheader("Platinum")
    st.write("- Unlimited uploaded files.")
    st.write("- Premium features and priority support.")

def main_page(username):
    # Header with Logo and Logout
    col1, col2, col3 = st.columns([0.2, 0.6, 0.2])
    with col1:
        st.image("images/logo.jpg", width=100)  # Replace "your_logo.png" with your logo file
    with col3:
        if username:
            if st.button("Logout"):
                st.query_params["page"] = "login"
                st.rerun()

    st.markdown("""
    <style>
    @keyframes fadeIn {
        from {opacity: 0;}
        to {opacity: 1;}
    }
    .stApp {
        animation: fadeIn 1.5s ease-in-out;
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([0.6, 0.4])

    with col1:
        st.header("Ask your PDFs, DOCX, Images, CSV, or Excel Files 💬")
        pdfs_and_docs = st.file_uploader("Upload PDFs, DOCX, Images (JPEG, PNG), CSV, or Excel", type=["pdf", "docx", "csv", "xlsx", "jpg", "jpeg", "png"], accept_multiple_files=True, label_visibility='collapsed')
        all_contacts = {"phone_numbers": set(), "emails": set()}

        user_data = load_user_data()  # Load it here
        if pdfs_and_docs:
            user_subscription = user_data.get(username, {}).get('subscription', 'free')
            uploaded_files = user_data.get(username, {}).get('uploaded_files', [])

            subscription_limits = {
                'free': 15,
                'silver': 50,
                'gold': 150,
                'platinum': float('inf')
            }

            if user_subscription == 'free' and len(uploaded_files) >= subscription_limits['free']:
                st.error(f"Your Free trial has exceeded the limit of {subscription_limits['free']} files. Please upgrade your subscription.")
                return

            user_subscription = user_subscription.capitalize() # Capitalize user_subscription
            if len(uploaded_files) + len(pdfs_and_docs) > subscription_limits[user_subscription.lower()]:
                st.error(f"Your current subscription ({user_subscription}) allows a maximum of {subscription_limits[user_subscription.lower()]} files. Please upgrade your subscription.")
                return

            all_chunks = []
            text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=200, length_function=len)

            for file in pdfs_and_docs:
                file_name = file.name
                prefixed_file_name = f"{username}_{file_name}"  # prefix the file name.

                if prefixed_file_name not in uploaded_files:
                    uploaded_files.append(prefixed_file_name)  # save the prefixed name to the user data

                file_extension = file_name.split(".")[-1].lower()

                file_contents = file.read()
                file_buffer = BytesIO(file_contents)
                file_buffer.seek(0)
                
                # Save file permanently
                save_uploaded_file(file, username)

                text = ""
                if file_extension == "pdf":
                    pdf_reader = PdfReader(BytesIO(file_contents))
                    extracted_text = [page.extract_text() for page in pdf_reader.pages if page.extract_text()]
                    text = "\n".join(extracted_text) if extracted_text else ""
                elif file_extension == "csv":
                    df = pd.read_csv(BytesIO(file_contents))
                    text = df.to_string(index=False)
                elif file_extension == "xlsx":
                    df = pd.read_excel(BytesIO(file_contents))
                    text = df.to_string(index=False)
                elif file_extension == "docx":
                    text = extract_text_from_docx(BytesIO(file_contents))
                elif file_extension in ["jpg", "jpeg", "png"]:
                    text = extract_text_from_image(BytesIO(file_contents))
                    logging.info(f"Extracted Text from {file_name}: {text[:100]}...")
                phone_numbers, emails = extract_contacts(text)
                all_contacts["phone_numbers"].update(phone_numbers)
                all_contacts["emails"].update(emails)
                chunks = text_splitter.split_text(text)
                logging.info(f"Chunks from {file_name}: {len(chunks)}")
                all_chunks.extend(chunks)

            logging.info(f"Total chunks before storing embeddings: {len(all_chunks)}")
            store_embeddings_permanently(all_chunks, username)

            user_data[username]['uploaded_files'] = uploaded_files
            save_user_data(user_data)

        index, documents = load_stored_embeddings(username)
        if index is None or documents is None:
            st.warning("No stored embeddings found. Upload files to build the knowledge base.")
            return

        user_question = st.text_input("Ask a question about your files:")
        if user_question:
            model = get_embedding_model()
            if model is None:
                st.error("Failed to load embedding model. Please try again.")
                return
                
            query_embedding = model.encode([user_question])[0]
            
            # Search using FAISS
            k = 5  # number of nearest neighbors
            distances, indices = index.search(query_embedding.reshape(1, -1).astype('float32'), k)
            
            # Filter results based on similarity threshold - more lenient threshold
            relevant_docs = []
            for idx, distance in zip(indices[0], distances[0]):
                if distance < 1.0:  # Increased threshold from 0.6 to 1.0
                    relevant_docs.append(documents[idx])
            
            if not relevant_docs:
                st.warning("No exact matches found. Here are the closest matches:")
                # If no documents pass the threshold, use the top 2 closest matches
                for idx, distance in zip(indices[0][:2], distances[0][:2]):
                    relevant_docs.append(documents[idx])
            
            if relevant_docs:
                try:
                    llm = OpenAI(openai_api_key=OPENAI_API_KEY)
                    chain = load_qa_chain(llm, chain_type="stuff")
                    with get_openai_callback() as cb:
                        response = chain.run(input_documents=relevant_docs, question=user_question)
                        logging.info(cb)
                        st.write(response)
                        # Copy to clipboard button
                        if st.button("Copy Response"):
                            st.session_state.copied_response = response
                            st.info("Response copied to clipboard!")
                except Exception as e:
                    st.error(f"Error generating response: {str(e)}")
                    if "authentication" in str(e).lower():
                        st.error("Please check your OpenAI API key in the .streamlit/secrets.toml file")
            else:
                st.error("No relevant information found in the uploaded files. Please try rephrasing your question or upload more relevant documents.")

    with col2:
        st.subheader("Uploaded Files")
        user_data = load_user_data()
        uploaded_files = user_data.get(username, {}).get('uploaded_files', [])
        if uploaded_files:
            with st.expander("📂 Click to view uploaded files"):
                for file_name in uploaded_files:
                    if file_name.startswith(f"{username}_"):
                        display_file_name = file_name[len(f"{username}_"):]
                        cols = st.columns([0.9, 0.1])
                        with cols[0]:
                            st.markdown(f"- 📄 **{display_file_name}**")
                        with cols[1]:
                            if st.button("❌", key=f"delete_{file_name}"):
                                delete_uploaded_file(file_name)
                                uploaded_files.remove(file_name)
                                user_data[username]['uploaded_files'] = uploaded_files
                                save_user_data(user_data)
                                st.rerun()
            if st.button("Delete All Files"):
                delete_all_uploaded_files(username)
                user_data[username]['uploaded_files'] = []
                save_user_data(user_data)
                st.rerun()
        else:
            st.info("No previously uploaded files.")

    # Navigation bar with Upgrade Subscription dropdown
    if username and username != "admin":
        pass

def main():
    page = st.query_params.get("page", "login")
    username = st.query_params.get("username", None)

    if page == "login":
        login()
    elif page == "signup":
        signup()
    elif page == "main":
        main_page(username)
        if username == "admin":
            admin_page(username)
    elif page == "admin":
        admin_page(username)

if __name__ == '__main__':
    main()